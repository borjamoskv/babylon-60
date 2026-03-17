from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, color_hex):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def configure_margins(doc):
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)


def configure_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Inter"
    font.size = Pt(11)


def add_cover_page(doc):
    doc.add_heading("", level=0)
    title = doc.add_heading("CORTEX-PERSIST", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Infraestructura de Confianza Descentralizada para Agentes de IA Autónomos"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(46, 80, 144)

    doc.add_paragraph("\n\n\n\n")
    author = doc.add_paragraph(
        "Análisis Técnico Exhaustivo\nAutor: Borja Fernández Angulo\nCORTEX System v4.0"
    )
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def add_table_of_contents(doc, section_titles):
    doc.add_heading("Tabla de Contenidos", level=1)
    for sec in section_titles:
        p = doc.add_paragraph(sec)
        p.style = doc.styles["List Number"]
    doc.add_page_break()


def add_section(doc, title, content):
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)


def add_table(doc, headers, data, col_widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for i, width in enumerate(col_widths):
        table.columns[i].width = Inches(width)

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_background(hdr_cells[i], "2E5090")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for row_data in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = item


def create_document():
    doc = Document()
    configure_margins(doc)
    configure_styles(doc)
    add_cover_page(doc)

    section_titles = [
        "1. El Imperativo Estratégico: La Crisis de Confianza",
        "2. Contexto Regulatorio: Cumplimiento por Diseño y EU AI Act",
        "3. El Motor de Confianza CORTEX (arquitectura de 5 capas)",
        "4. Sistema de Memoria Cognitiva Tripartita (L1/L2/L3)",
        "5. Criptografía y Verificación (SHA-256, Merkle Trees, Sandbox AST)",
        "6. Consenso Multi-Agente (WBFT)",
        "7. Interoperabilidad y Despliegue Multiplataforma",
        "8. Metrología del Código Base (45,500 LOC, 444 módulos)",
        "9. Conclusiones Estratégicas",
    ]

    add_table_of_contents(doc, section_titles)

    add_section(
        doc,
        section_titles[0],
        "La adopción de agentes de inteligencia artificial autónomos en entornos empresariales críticos se ve frenada por una crisis fundamental: la falta de confianza. Las organizaciones requieren garantías irrefutables de que los agentes operan dentro de límites predefinidos, que sus acciones son trazables y que la infraestructura subyacente puede resistir manipulaciones maliciosas o derivas estocásticas. CORTEX-Persist emerge como la solución arquitectónica a esta anomalía estructurada.",
    )

    add_section(
        doc,
        section_titles[1],
        "CORTEX implementa un modelo de 'Cumplimiento por Diseño', asegurando que la arquitectura base satisface nativamente los requisitos regulatorios globales, en particular la EU AI Act.",
    )
    add_table(
        doc,
        ["Requisito EU AI Act", "Implementación en CORTEX"],
        [
            (
                "Transparencia y Trazabilidad",
                "Log inmutable criptográfico de decisiones y acciones de los agentes (Chain of Thought Ledger).",
            ),
            (
                "Supervisión Humana",
                'Protocolos "Human-in-the-loop" obligatorios en operaciones críticas (C-Level Sandbox).',
            ),
            (
                "Gestión de Riesgos",
                "Evaluación estocástica y sandboxing estricto de ejecución de código (AST validation).",
            ),
        ],
        [2.5, 4.0],
    )

    add_section(
        doc,
        section_titles[2],
        "El núcleo de CORTEX es su motor de confianza, segmentado en cinco bi-capas estructurales para garantizar aislamiento absoluto y comunicación segura (Zero-Trust).",
    )
    add_table(
        doc,
        ["Capa", "Responsabilidad e Interfaz"],
        [
            ("L0 - Base", "Sistema subyacente y sistema de archivos cifrado (AES-256)."),
            ("L1 - Engine", "Motor de ejecución determinista y Sandbox AST."),
            ("L2 - Memory", "Memoria transaccional L1/L2/L3 (Redis/SQLite)."),
            ("L3 - Agent", "Máquina de estados finitos del Agente y WBFT."),
            ("L4 - Network", "Protocolos M2M cifrados y RPC sobre TLS 1.3."),
        ],
        [1.5, 5.0],
    )

    add_section(
        doc,
        section_titles[3],
        "El sistema de memoria simula la cognición mamífera avanzada, permitiendo a los agentes razonar temporalmente sin colapso de contexto.",
    )
    add_table(
        doc,
        ["Nivel Cognitivo", "Tecnología", "Persistencia / TTL"],
        [
            ("L1: Working Memory", "Redis (In-Memory)", "Efímera (Volátil, sub-milisegundo)."),
            (
                "L2: Episodic Memory",
                "SQLiteDB (Transaccional)",
                "Sesión / Proyecto (Indices semánticos).",
            ),
            ("L3: Semantic Core", "Vector DB", "Permanente (Conocimiento ontológico)."),
        ],
        [2.1, 2.1, 2.1],
    )

    add_section(
        doc,
        section_titles[4],
        "La integridad de los datos en CORTEX está garantizada a través de primitivas criptográficas robustas y un entorno de validación abstracto (AST). Cada modificación de estado genera un hash que se entrelaza en un Árbol de Merkle.",
    )

    add_section(
        doc,
        section_titles[5],
        "La comunicación en enjambre requiere tolerancia a fallos. CORTEX implementa una variante de Tolerancia a Faltas Bizantinas Ponderada (WBFT).",
    )
    add_table(
        doc,
        ["Mecanismo", "Implementación CORTEX-WBFT"],
        [
            (
                "Leader Election",
                "Rotación determinista basada en reputación histórica (proof-of-accuracy).",
            ),
            ("Quorum", "Super-mayoría 2/3 para decisiones de Nivel 3 (Modificaciones de código)."),
            (
                "Slashing",
                "Penalización de tokens operacionales para agentes divergentes o maliciosos.",
            ),
        ],
        [2.0, 4.5],
    )

    add_section(
        doc,
        section_titles[6],
        "La arquitectura de CORTEX soporta portabilidad nativa y resiliencia de red.",
    )
    add_table(
        doc,
        ["Entorno", "Técnica de Aislamiento"],
        [
            (
                "Bare-Metal (Desktop)",
                "Local Daemon con privilegios limitados y chroot (NotchLive).",
            ),
            ("Nube (K8s/Docker)", "Microservicios L1/L2/L3 orquestados mediante CORTEX Nexus."),
            ("Edge/IoT", "CORTEX-Lite (Rust footprint de baja latencia)."),
        ],
        [2.0, 4.5],
    )

    add_section(
        doc,
        section_titles[7],
        "Para mantener la robustez sin caer en el gigantismo de software, CORTEX restringe su código base enfocándose en la densidad y el determinismo.",
    )
    add_table(
        doc,
        ["Métrica (CORTEX v4)", "Valor"],
        [
            ("Líneas de Código (LOC)", "45,500 (Densidad optimizada)"),
            ("Módulos Estructurales", "444 componentes aislados"),
            ("Cobertura de Tests", "> 94% (Unit, E2E, Chaos)"),
            ("Latencia Promedio L1", "< 15ms (P99)"),
        ],
        [2.5, 4.0],
    )

    add_section(
        doc,
        section_titles[8],
        "La adopción de CORTEX-Persist marca la transición de simples bots conversacionales de IA a flotas de agentes autónomos económicamente viables y legalmente defendibles. La asimetría entrópica resuelta por su arquitectura L1 a L3 posiciona a las organizaciones como inmunes a las ineficiencias de enjambres no estructurados y riesgos regulatorios masivos.",
    )

    doc.add_page_break()

    doc.add_heading("Contraportada", level=1)
    doc.add_paragraph(
        "Proyecto CORTEX-PERSIST\n"
        "Arquitectura Soberana MOSKV-1\n"
        "Licencia Privada (Strictly Confidential)\n\n"
        "© 2026 Borja Fernández Angulo"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        doc.save("Cortex_Persist_Analisis.docx")
        print("Documento guardado: Cortex_Persist_Analisis.docx")
    except Exception as e:
        print(f"Error saving document: {e}")


if __name__ == "__main__":
    create_document()
